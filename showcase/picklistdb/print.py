import sys
import os
import win32print
import subprocess
import pythoncom
from win32com import client
import time
import barcode
from docx import Document
from docx.shared import Cm
# from docx.text.run import WD_BREAK

# THIS IS NOT THREAD SAFE. IF MULTIPLE PRINTS ARE CALLED AT THE SAME TIME, EXPECT ISSUES.
def printWordDocument(filename):
    word = client.Dispatch("Word.Application")
    word.Documents.Open(filename)
    word.ActiveDocument.PrintOut()
    time.sleep(2)
    word.ActiveDocument.Close()
    word.Quit()


def print_barcodes(list_of_ids, new_cwd):
    listy = []
    os.chdir(new_cwd)
    print(os.getcwd())
    pythoncom.CoInitialize()
    from barcode.writer import ImageWriter
    for items in list_of_ids:
        EAN = barcode.get_barcode_class('code128')
        if not os.path.exists(os.getcwd()+"/barcode tempy"):
            os.mkdir(os.getcwd()+"/barcode tempy")
        print(items)
        listy.append(EAN(items, writer=ImageWriter()))

    document = Document()  # generate a new temporary document.
    document_target = os.getcwd() + "/tempy.docx"
    # changing the page margins
    sections = document.sections
    for i in range(len(listy)):
        listy[i].save(os.getcwd()+"/barcode tempy/" + str(i))
        print("saved @ " + os.getcwd() + "/barcode tempy/" + str(i))
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
    p = document.add_paragraph()  # add a new paragraph
    r = p.add_run()
    flag = False
    for i in range(0, len(listy)):
        r.add_picture(os.getcwd()+"/barcode tempy/"+str(i)+".png", width=Cm(10.4), height=Cm(4.9))
        if flag:
            r.add_break()
            r.add_break()
            if i % 3 == 0:
                r.add_break()
                r.add_break()
        if len(listy)-i > 2:
            flag = not flag
    document.save(document_target)

    printWordDocument(document_target)


